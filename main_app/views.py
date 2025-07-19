import pandas as pd
import re
import os
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import mammoth
import time
from datetime import datetime
from io import BytesIO

from django.shortcuts import render, redirect
from django.http import FileResponse, HttpResponse, Http404
from django.contrib import messages
from django.conf import settings
from django.utils import timezone

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.utils import get_column_letter
import json

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from .forms import ExcelUploadForm, MailingListUploadForm, SmtpCredentialsForm


def clean_header(s):
    if not isinstance(s, str):
        return ''
    s = re.sub(r'-\s*\n\s*', '', s)
    s = re.sub(r'[\r\n]+', ' ', s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def is_subheader(row):
    values = row.dropna().unique()
    return len(values) == 1


def find_col(columns, keywords):
    for col in columns:
        cleaned_col = str(col).replace('\n', '').replace('\r', '').strip()
        for kw in keywords:
            if kw.lower() in cleaned_col.lower():
                return col
    return None


def save_formatted_excel(df, filename, project_name, estimate_section):
    wb = Workbook()
    ws = wb.active
    ws.title = "Отфильтрованные данные"

    header_font = Font(name='Arial', size=11, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    info_font = Font(name='Arial', size=10, bold=True)
    data_font = Font(name='Arial', size=10)
    data_alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
    border_thin = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    ws['A1'] = "Наименование объекта:"
    ws['B1'] = project_name
    ws['A2'] = "Раздел сметы:"
    ws['B2'] = estimate_section

    for row_num in [1, 2]:
        ws[f'A{row_num}'].font = info_font
        ws[f'B{row_num}'].font = Font(name='Arial', size=10)

    start_row = 4

    for col_idx, header in enumerate(df.columns, 1):
        cell = ws.cell(row=start_row, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = border_thin

    for row_idx, row in enumerate(dataframe_to_rows(df, index=False, header=False), start_row + 1):
        for col_idx, value in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = data_font
            cell.alignment = data_alignment
            cell.border = border_thin
            if (row_idx - start_row) % 2 == 0:
                cell.fill = PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid')

    for col_idx in range(1, len(df.columns) + 1):
        max_length = 0
        column_letter = get_column_letter(col_idx)
        header_length = len(str(df.columns[col_idx - 1]))
        max_length = max(max_length, header_length)

        for row_idx in range(start_row + 1, ws.max_row + 1):
            cell_value = ws.cell(row=row_idx, column=col_idx).value
            if cell_value is not None:
                cell_length = len(str(cell_value))
                max_length = max(max_length, cell_length)

        adjusted_width = min(max(max_length + 2, 12), 50)
        ws.column_dimensions[column_letter].width = adjusted_width

    ws.freeze_panes = f'A{start_row + 1}'
    ws.row_dimensions[1].height = 20
    ws.row_dimensions[2].height = 20
    ws.row_dimensions[start_row].height = 30

    wb.save(filename)


def get_main_letter_text(company_name, project_name, delivery_address, project_docs_link="", include_table_text=True):
    if not delivery_address:
        delivery_address = project_name

    main_text = f"""Добрый день!

{company_name} просит Вас предоставить коммерческое предложение (счет). В рамках реализации проекта {project_name}. Цены интересуют оптовые с максимальной скидкой.

Просьба, стоимость доставки включить в стоимость материалов и оборудования. Прописать: Базис поставки: {delivery_address}.

Также коммерческое предложение должно содержать информацию о ставке НДС, номере и дате коммерческого предложения, сроке действия предложения, реквизиты Вашей организации, информацию об учете дополнительных услуг, таких как: шефмонтаж, шефналадка. Печать, подпись уполномоченного лица."""

    if project_docs_link:
        main_text += f"\n\nПроектная документация доступна по ссылке: {project_docs_link}."

    if include_table_text:
        main_text += """

Перечень необходимых материалов в приложенной таблице ниже."""

    main_text += """

При возникновении вопросов, просьба звонить по номеру телефона, указанному в подписи данного сообщения. Или писать на этот адрес почты, ответным письмом."""

    return main_text


def create_commercial_request_from_template(df_filtered, project_name, estimate_section,
                                            template_path="template_commercial_request.docx",
                                            company_name="[НАЗВАНИЕ КОМПАНИИ]",
                                            delivery_address="",
                                            company_details="[РЕКВИЗИТЫ КОМПАНИИ]",
                                            contact_person="[КОНТАКТНОЕ ЛИЦО И НОМЕР]",
                                            project_docs_link=""):
    try:
        doc = Document(template_path)
    except Exception as e:
        messages.warning("Не удалось загрузить шаблон DOCX. Создан базовый документ.")
        return create_basic_template_document(df_filtered, project_name, estimate_section,
                                              company_name, delivery_address,
                                              company_details, contact_person, project_docs_link)

    if not delivery_address:
        delivery_address = project_name
    df_for_table = df_filtered.copy()
    columns_to_exclude = ['Обоснование']

    for col in columns_to_exclude:
        if col in df_for_table.columns:
            df_for_table = df_for_table.drop(columns=[col])

    should_include_table = len(df_for_table) <= 10
    main_text = get_main_letter_text(company_name, project_name, delivery_address,
                                     project_docs_link, include_table_text=should_include_table)

    replacements = {
        '{COMPANY_NAME}': company_name,
        '{PROJECT_NAME}': project_name,
        '{DELIVERY_ADDRESS}': delivery_address,
        '{ESTIMATE_SECTION}': estimate_section,
        '{COMPANY_DETAILS}': company_details,
        '{CONTACT_PERSON}': contact_person,
        '{PROJECT_DOCS_LINK}': project_docs_link,
        '{MAIN_TEXT}': main_text
    }

    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph.text:
                replace_text_in_paragraph(paragraph, placeholder, replacement)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in paragraph.text:
                            replace_text_in_paragraph(paragraph, placeholder, replacement)

    table_inserted = False
    for paragraph in doc.paragraphs:
        if '{TABLE_PLACEHOLDER}' in paragraph.text:
            if should_include_table:
                p_element = paragraph._element
                parent = p_element.getparent()

                paragraph.clear()

                paragraph.add_run("ПЕРЕЧЕНЬ НЕОБХОДИМЫХ МАТЕРИАЛОВ:").bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                table = doc.add_table(rows=1, cols=len(df_for_table.columns))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                table_element = table._tbl
                parent.insert(parent.index(p_element) + 1, table_element)

                header_cells = table.rows[0].cells
                for j, column_name in enumerate(df_for_table.columns):
                    header_cells[j].text = str(column_name)
                    for paragraph_cell in header_cells[j].paragraphs:
                        for run in paragraph_cell.runs:
                            run.bold = True
                    header_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for index, row in df_for_table.iterrows():
                    row_cells = table.add_row().cells
                    for j, value in enumerate(row):
                        row_cells[j].text = str(value) if pd.notna(value) else ''

                for j, column in enumerate(table.columns):
                    for cell in column.cells:
                        cell.width = Inches(2.0)

                table_inserted = True
                paragraph.add_run(
                    "Перечень необходимых материалов Вы можете также найти в прилагаемом файле.").italic = True
            else:
                paragraph.clear()
                paragraph.add_run("Перечень необходимых материалов Вы можете найти в прилагаемом файле.").italic = True
                table_inserted = True
            break

    if not table_inserted:
        if should_include_table:
            add_materials_table_to_end(doc, df_for_table)

    return doc


def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        alignment = paragraph.alignment

        full_text = paragraph.text

        if ':' in full_text and old_text in full_text:
            parts = full_text.split(':', 1)
            if len(parts) == 2:
                before_colon = parts[0] + ':'
                after_colon = parts[1]

                if old_text in after_colon:
                    paragraph.clear()
                    bold_run = paragraph.add_run(before_colon)
                    bold_run.bold = True
                    normal_text = after_colon.replace(old_text, new_text)
                    normal_run = paragraph.add_run(normal_text)
                    normal_run.bold = False
                    if alignment:
                        paragraph.alignment = alignment
                    return

        first_run = paragraph.runs[0] if paragraph.runs else None
        paragraph.clear()
        new_text_content = full_text.replace(old_text, new_text)
        new_run = paragraph.add_run(new_text_content)

        if first_run:
            new_run.bold = first_run.bold
            new_run.italic = first_run.italic
            new_run.underline = first_run.underline
            if first_run.font.size:
                new_run.font.size = first_run.font.size
            if first_run.font.name:
                new_run.font.name = first_run.font.name

        if alignment:
            paragraph.alignment = alignment


def add_materials_table_to_end(doc, df_for_table):
    doc.add_paragraph()
    materials_header = doc.add_paragraph()
    materials_header.add_run("ПЕРЕЧЕНЬ НЕОБХОДИМЫХ МАТЕРИАЛОВ:").bold = True
    materials_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=len(df_for_table.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, column_name in enumerate(df_for_table.columns):
        header_cells[i].text = str(column_name)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for index, row in df_for_table.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value) if pd.notna(value) else ''

    for i, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = Inches(2.0)


def create_basic_template_document(df_filtered, project_name, estimate_section,
                                   company_name, delivery_address,
                                   company_details, contact_person, project_docs_link=""):
    doc = Document()

    title = doc.add_heading('ЗАПРОС КОММЕРЧЕСКОГО ПРЕДЛОЖЕНИЯ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    if not delivery_address:
        delivery_address = project_name

    df_for_table = df_filtered.copy()
    columns_to_exclude = ['Обоснование']

    for col in columns_to_exclude:
        if col in df_for_table.columns:
            df_for_table = df_for_table.drop(columns=[col])

    should_include_table = len(df_for_table) <= 10

    main_text = get_main_letter_text(company_name, project_name, delivery_address,
                                     project_docs_link, include_table_text=should_include_table)
    doc.add_paragraph(main_text)

    doc.add_paragraph()
    info_paragraph = doc.add_paragraph()
    info_paragraph.add_run("Информация о проекте:").bold = True

    project_info = doc.add_paragraph()
    project_info.add_run("• Наименование объекта: ").bold = True
    project_info.add_run(project_name).bold = False

    section_info = doc.add_paragraph()
    section_info.add_run("• Раздел сметы: ").bold = True
    section_info.add_run(estimate_section).bold = False

    if should_include_table:
        doc.add_paragraph()
        materials_header = doc.add_paragraph()
        materials_header.add_run("ПЕРЕЧЕНЬ НЕОБХОДИМЫХ МАТЕРИАЛОВ:").bold = True
        materials_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=1, cols=len(df_for_table.columns))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = table.rows[0].cells
        for i, column_name in enumerate(df_for_table.columns):
            header_cells[i].text = str(column_name)
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for index, row in df_for_table.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ''

        for i, column in enumerate(table.columns):
            for cell in column.cells:
                cell.width = Inches(2.0)

        doc.add_paragraph()
        note_paragraph = doc.add_paragraph()
        note_paragraph.add_run(
            "Перечень необходимых материалов Вы можете также найти в прилагаемом файле.").italic = True
    else:
        doc.add_paragraph()
        note_paragraph = doc.add_paragraph()
        note_paragraph.add_run("Перечень необходимых материалов Вы можете найти в прилагаемом файле.").italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    requisites_header = doc.add_paragraph()
    requisites_header.add_run("Реквизиты компании:").bold = True

    requisites_text = doc.add_paragraph(company_details)

    doc.add_paragraph()
    contact_paragraph = doc.add_paragraph(contact_person)

    return doc


def create_template_file():
    doc = Document()
    title = doc.add_heading('ЗАПРОС КОММЕРЧЕСКОГО ПРЕДЛОЖЕНИЯ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    doc.add_paragraph("{MAIN_TEXT}")

    doc.add_paragraph()
    info_paragraph = doc.add_paragraph()
    info_paragraph.add_run("Информация о проекте:").bold = True

    project_info = doc.add_paragraph()
    project_info.add_run("• Наименование объекта: ").bold = True
    project_info.add_run("{PROJECT_NAME}")

    section_info = doc.add_paragraph()
    section_info.add_run("• Раздел сметы: ").bold = True
    section_info.add_run("{ESTIMATE_SECTION}")

    doc.add_paragraph()
    table_placeholder = doc.add_paragraph("{TABLE_PLACEHOLDER}")
    doc.add_paragraph()
    doc.add_paragraph()

    requisites_header = doc.add_paragraph()
    requisites_header.add_run("Реквизиты компании:").bold = True

    requisites_text = doc.add_paragraph("{COMPANY_DETAILS}")

    doc.add_paragraph()
    contact_paragraph = doc.add_paragraph("{CONTACT_PERSON}")

    return doc


def is_valid_email(email):
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email.strip()) is not None


def parse_emails(email_string):
    if pd.isna(email_string) or email_string == '' or str(email_string).lower() == 'nan':
        return []
    emails = [email.strip() for email in str(email_string).split(';')]
    valid_emails = [email for email in emails if is_valid_email(email)]
    return valid_emails


def extract_html_from_word(file_content_bytes):
    try:
        with BytesIO(file_content_bytes) as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
        return html
    except Exception as e:
        print(f"❌ Ошибка при извлечении HTML из Word: {str(e)}")
        return None


def replace_placeholders_in_html(html_template, company_data):
    if not html_template:
        return ""

    text = html_template

    replacements = {
        '{наименование_компании}': str(company_data.get('Наименование компании', '')),
        '{юридический_адрес}': str(company_data.get('Юридический адрес', '')),
        '{инн}': str(company_data.get('ИНН', '')),
        '{ссылка_на_сайт}': str(company_data.get('Ссылка на сайт', '')),
        '{контактный_телефон}': str(company_data.get('Контактный телефон', '')),
        '{электронная_почта}': str(company_data.get('Электронная почта', '')),
        '{город}': str(company_data.get('Город', '')),
        '{регион}': str(company_data.get('Регион', ''))
    }

    for placeholder, value in replacements.items():
        text = text.replace(placeholder, str(value) if pd.notna(value) and str(value).lower() != 'nan' else "")

    return text


def send_email_batch(smtp_config, from_email, to_emails, subject, body, attachment_paths=None):
    try:
        msg = MIMEMultipart()
        msg['From'] = from_email
        msg['To'] = ', '.join(to_emails)
        msg['Subject'] = subject

        msg.attach(MIMEText(body, 'html', 'utf-8'))

        if attachment_paths:
            for attachment_path in attachment_paths:
                if os.path.exists(attachment_path):
                    with open(attachment_path, 'rb') as attachment:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(attachment.read())
                    encoders.encode_base64(part)
                    filename = os.path.basename(attachment_path)
                    part.add_header(
                        'Content-Disposition',
                        f'attachment; filename="{filename}"'
                    )
                    msg.attach(part)
                else:
                    print(f"Attachment not found: {attachment_path}")

        with smtplib.SMTP_SSL(smtp_config['server'], smtp_config['port']) as server:
            server.login(smtp_config['user'], smtp_config['password'])
            server.send_message(msg)

        return True, None

    except Exception as e:
        return False, str(e)


def process_excel_file(excel_file):
    df_raw = pd.read_excel(excel_file, header=None)

    project_name = df_raw.iloc[1].dropna().values[0].strip() if not df_raw.iloc[
        1].dropna().empty else "Не определено"
    raw_estimate_section = df_raw.iloc[6].dropna().values[0].strip() if not df_raw.iloc[
        6].dropna().empty else "Не определено"
    estimate_section = raw_estimate_section if raw_estimate_section != "Раздел" else "Не определено"

    header_row_1 = df_raw.iloc[9].fillna('')
    header_row_2 = df_raw.iloc[10].fillna('')

    header_row_1_clean = header_row_1.apply(clean_header)
    header_row_2_clean = header_row_2.apply(clean_header)

    table_headers = [
        ' '.join(filter(None, [h1, h2])).strip()
        for h1, h2 in zip(header_row_1_clean, header_row_2_clean)
    ]

    df_data = df_raw.iloc[11:].copy()
    df_data.columns = table_headers

    df_data = df_data[~df_data.apply(is_subheader, axis=1)]

    col_obosn = find_col(df_data.columns, ["обоснование"])
    col_name = find_col(df_data.columns, ["наименование"])
    col_unit = find_col(df_data.columns, ["ед", "изм"])
    col_qty = find_col(df_data.columns, ["общее кол", "кол-во"])
    col_pp = find_col(df_data.columns, ["№ пп"])

    missing_cols = []
    if not col_obosn: missing_cols.append("Обоснование")
    if not col_name: missing_cols.append("Наименование")
    if not col_unit: missing_cols.append("Ед. изм.")
    if not col_qty: missing_cols.append("Количество")

    if missing_cols:
        raise ValueError(
            f"Не найдены обязательные столбцы: {', '.join(missing_cols)}. Проверьте названия и структуру файла.")

    mask = df_data[col_obosn].astype(str).str.startswith(('ТЦ', 'ФССЦ'), na=False)

    df_filtered = df_data.loc[mask, [col_obosn, col_name, col_unit, col_qty]].copy()

    df_filtered.rename(columns={
        col_obosn: 'Обоснование',
        col_name: 'Наименование',
        col_unit: 'Ед. изм.',
        col_qty: 'Количество',
    }, inplace=True)

    if col_pp:
        temp_pp_series = df_data.loc[mask, col_pp].fillna(method='ffill')
        df_filtered.insert(0, '№ по порядку', temp_pp_series.reset_index(drop=True))

    output_dir = os.path.join(settings.MEDIA_ROOT, 'processed_excels')
    os.makedirs(output_dir, exist_ok=True)

    timestamp = timezone.now().strftime('%Y%m%d%H%M%S')
    file_name_original, _ = os.path.splitext(excel_file.name)
    output_filename = f"request-for-prices_{file_name_original}_{timestamp}.xlsx"
    output_file_path = os.path.join(output_dir, output_filename)

    save_formatted_excel(df_filtered, output_file_path, project_name, estimate_section)

    return output_file_path, project_name, estimate_section, df_filtered

def generate_docx_letter_from_excel_data(request, df_filtered, project_name, estimate_section):
    template_filename = "template_commercial_request.docx"
    template_docx_path = os.path.join(settings.BASE_DIR, 'main_app', 'templates', 'main_app', template_filename)

    if not os.path.exists(template_docx_path):
        template_docx_path = os.path.join(settings.BASE_DIR, 'main_app', 'templates', template_filename)
        if not os.path.exists(template_docx_path):
            # Now pass 'request' here
            messages.warning(request, "Шаблон DOCX не найден. Создается базовый шаблон письма.")
            base_template_doc = create_template_file()
            base_template_path = os.path.join(settings.MEDIA_ROOT, 'generated_templates', 'base_commercial_request_template.docx')
            os.makedirs(os.path.dirname(base_template_path), exist_ok=True)
            base_template_doc.save(base_template_path)
            template_docx_path = base_template_path

    COMPANY_SETTINGS = {
        'company_name': 'ООО "СМР"',
        'delivery_address': '',
        'company_details': """Общество с ограниченной ответственностью "СТРОЙМОНТАЖРЕКОНСТРУКЦИЯ"
ИНН: 7500009942
КПП: 750001001
Адрес: 672012, Чита, ул.Нагорная, д.43, пом.6, каб.3
Банк: Читинское РФ АО Россельхозбанк
Р/С: 40702810447000001638
К/С: 30101810400000000740
БИК: 047601740""",
        'contact_person': 'Ольга, тел.: +7 (919) 901-45-47',
        'project_docs_link': 'https://cloud.smr-company.ru/projects/docs/test-project-2024'
    }

    doc = create_commercial_request_from_template(
        df_filtered=df_filtered,
        project_name=project_name,
        estimate_section=estimate_section,
        template_path=template_docx_path,
        **COMPANY_SETTINGS
    )

    docx_output_dir = os.path.join(settings.MEDIA_ROOT, 'generated_docs')
    os.makedirs(docx_output_dir, exist_ok=True)

    cleaned_project_name_for_save = re.sub(r'[^\w\-. ]', '',
                                           project_name.replace('\n', ' ').replace('\r', ' ').strip())
    docx_output_filename = f"Коммерческий_запрос_{cleaned_project_name_for_save}_{timezone.now().strftime('%Y%m%d%H%M%S')}.docx"
    docx_output_path = os.path.join(docx_output_dir, docx_output_filename)

    doc.save(docx_output_path)
    return docx_output_path

def save_csv_for_mailing(csv_file):
    output_dir = os.path.join(settings.MEDIA_ROOT, 'mailing_lists')
    os.makedirs(output_dir, exist_ok=True)

    timestamp = timezone.now().strftime('%Y%m%d%H%M%S')
    file_name_original, _ = os.path.splitext(csv_file.name)
    output_filename = f"mailing_list_{file_name_original}_{timestamp}.csv"
    output_file_path = os.path.join(output_dir, output_filename)

    with open(output_file_path, 'wb+') as destination:
        for chunk in csv_file.chunks():
            destination.write(chunk)

    return output_file_path


def download_processed_excel(request):
    file_path = request.GET.get('filepath')
    if not file_path:
        raise Http404("Путь к файлу не указан.")

    allowed_dir = os.path.join(settings.MEDIA_ROOT, 'processed_excels')
    if not os.path.exists(file_path) or not os.path.isfile(file_path) or \
            not os.path.commonpath([os.path.realpath(file_path), os.path.realpath(allowed_dir)]) == os.path.realpath(
                allowed_dir):
        raise Http404("Файл не найден или указан недопустимый путь.")

    if not file_path.lower().endswith('.xlsx'):
        raise Http404("Недопустимый тип файла.")

    try:
        download_filename = os.path.basename(file_path)
        if download_filename.startswith("request-for-prices_"):
            download_filename = download_filename[len("request-for-prices_"):]
        download_filename = f"Отфильтрованный_прайс_{download_filename}"

        response = FileResponse(open(file_path, 'rb'), as_attachment=True, filename=download_filename)
        return response
    except Exception as e:
        return HttpResponse(f"Ошибка при обслуживании файла: {e}", status=500)


def upload_excel(request):
    excel_form = ExcelUploadForm()
    mailing_list_form = MailingListUploadForm()

    processed_excel_url = request.session.get('processed_excel_url', None)
    project_name = request.session.get('project_name', 'Не определено')
    estimate_section = request.session.get('estimate_section', 'Не определено')
    processed_docx_path = request.session.get('processed_docx_path', None)
    last_uploaded_csv_path = request.session.get('last_uploaded_csv_path', None)

    if processed_excel_url and (not request.session.get('last_processed_excel_path') or not os.path.exists(
            request.session['last_processed_excel_path'])):
        messages.warning(request, "Предыдущий обработанный Excel-файл не найден или устарел. Загрузите новый.")
        del request.session['processed_excel_url']
        del request.session['project_name']
        del request.session['estimate_section']
        processed_excel_url = None
        project_name = "Не определено"
        estimate_section = "Не определено"

    if processed_docx_path and not os.path.exists(processed_docx_path):
        messages.warning(request, "Сгенерированное письмо DOCX не найдено. Пожалуйста, сгенерируйте его снова.")
        del request.session['processed_docx_path']
        processed_docx_path = None

    if last_uploaded_csv_path and not os.path.exists(last_uploaded_csv_path):
        messages.warning(request, "Загруженный список рассылки CSV не найден. Пожалуйста, загрузите его снова.")
        del request.session['last_uploaded_csv_path']
        last_uploaded_csv_path = None

    if request.method == 'POST':
        if 'excel_file' in request.FILES:
            excel_form = ExcelUploadForm(request.POST, request.FILES)
            if excel_form.is_valid():
                excel_file = request.FILES['excel_file']
                if not excel_file.name.lower().endswith('.xlsx'):
                    messages.error(request, "Пожалуйста, загрузите файл формата .xlsx.")
                else:
                    try:
                        output_file_path, proj_name, est_section, df_filtered_data = process_excel_file(excel_file)

                        request.session['last_processed_excel_path'] = output_file_path
                        request.session[
                            'processed_excel_url'] = f'/download-processed-excel/?filepath={output_file_path}'
                        request.session['project_name'] = proj_name
                        request.session['estimate_section'] = est_section
                        request.session['processed_docx_path'] = None
                        request.session['last_uploaded_csv_path'] = None

                        messages.success(request,
                                         "Excel файл успешно обработан! Теперь вы можете сгенерировать письмо или загрузить список рассылки.")
                        return redirect('upload_excel')
                    except ValueError as ve:
                        messages.error(request, f"Ошибка данных в Excel: {str(ve)}")
                    except Exception as e:
                        messages.error(request,
                                       f"Ошибка при обработке Excel-файла: {str(e)}. Проверьте структуру файла.")
            else:
                messages.error(request, "Ошибка в форме загрузки Excel. Пожалуйста, выберите файл Excel.")

        elif request.POST.get('action') == 'prepare_letters':
            if not processed_excel_url or not request.session.get('last_processed_excel_path') or \
                    not os.path.exists(request.session['last_processed_excel_path']):
                messages.error(request,
                               "Обработанный Excel-файл не найден или устарел. Пожалуйста, загрузите Excel снова (Шаг 1).")
                return redirect('upload_excel')

            try:
                df_filtered_for_docx = pd.read_excel(request.session['last_processed_excel_path'])

                docx_output_path = generate_docx_letter_from_excel_data(
                    request,
                    df_filtered=df_filtered_for_docx,
                    project_name=request.session['project_name'],
                    estimate_section=request.session['estimate_section']
                )

                request.session['processed_docx_path'] = docx_output_path
                request.session['last_uploaded_csv_path'] = None

                messages.success(request, 'DOCX письмо успешно сгенерировано!')
                return redirect('upload_excel')
            except FileNotFoundError as fnfe:
                messages.error(request, f"Ошибка: {str(fnfe)}")
            except Exception as e:
                messages.error(request, f"Ошибка при генерации DOCX письма: {str(e)}")

        elif 'csv_file' in request.FILES:
            mailing_list_form = MailingListUploadForm(request.POST, request.FILES)
            if mailing_list_form.is_valid():
                csv_file = request.FILES['csv_file']

                if not csv_file.name.lower().endswith('.csv'):
                    messages.error(request, "Пожалуйста, загрузите файл формата .csv.")
                else:
                    if not processed_docx_path or not os.path.exists(processed_docx_path):
                        messages.warning(request,
                                         'Пожалуйста, сгенерируйте DOCX письмо (Шаг 2), прежде чем загружать список рассылки.')
                    else:
                        try:
                            csv_output_path = save_csv_for_mailing(csv_file)
                            request.session['last_uploaded_csv_path'] = csv_output_path
                            messages.success(request, 'Список рассылки CSV успешно загружен!')
                            return redirect('upload_excel')
                        except Exception as e:
                            messages.error(request, f"Ошибка при загрузке списка рассылки CSV: {str(e)}")
            else:
                messages.error(request, "Ошибка в форме загрузки списка рассылки CSV. Пожалуйста, выберите файл CSV.")

    excel_form = ExcelUploadForm(request.POST or None, request.FILES or None)
    mailing_list_form = MailingListUploadForm(request.POST or None, request.FILES or None)

    context = {
        'excel_form': excel_form,
        'mailing_list_form': mailing_list_form,
        'processed_file_url': processed_excel_url,
        'project_name': project_name,
        'estimate_section': estimate_section,
        'processed_docx_path': processed_docx_path,
        'last_uploaded_csv_path': last_uploaded_csv_path,
    }
    return render(request, 'main_app/upload_excel.html', context)


def home(request):
    return render(request, 'main_app/home.html')

import os
import pandas as pd
from django.shortcuts import render, redirect
from django.contrib import messages
# from .forms import SmtpCredentialsForm
# from .utils import extract_html_from_word, parse_emails, send_email_batch
def is_valid_email(email):
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email.strip()) is not None

def parse_emails(email_string):
    if pd.isna(email_string) or email_string == '' or str(email_string).lower() == 'nan':
        return []
    emails = [email.strip() for email in str(email_string).split(';')]
    valid_emails = [email for email in emails if is_valid_email(email)]
    return valid_emails

def replace_placeholders(template_text, company_data):
    if not template_text:
        return ""
    text = template_text
    replacements = {
        '{наименование_компании}': str(company_data.get('Наименование компании', '')),
        '{юридический_адрес}': str(company_data.get('Юридический адрес', '')),
        '{инн}': str(company_data.get('ИНН', '')),
        '{ссылка_на_сайт}': str(company_data.get('Ссылка на сайт', '')),
        '{контактный_телефон}': str(company_data.get('Контактный телефон', '')),
        '{электронная_почта}': str(company_data.get('Электронная почта', '')),
        '{город}': str(company_data.get('Город', '')),
        '{регион}': str(company_data.get('Регион', ''))
    }
    for placeholder, value in replacements.items():
        text = text.replace(placeholder, value if pd.notna(value) and value != 'nan' else "")
    return text

def send_email_batch(smtp_config, from_email, to_emails, subject, body, attachment_paths=None):
    try:
        msg = MIMEMultipart()
        msg['From'] = from_email
        msg['To'] = ', '.join(to_emails)
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'html', 'utf-8'))

        if attachment_paths:
            for attachment_path in attachment_paths:
                if attachment_path and os.path.exists(attachment_path):
                    with open(attachment_path, 'rb') as attachment:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(attachment.read())
                    encoders.encode_base64(part)
                    filename = os.path.basename(attachment_path)
                    part.add_header(
                        'Content-Disposition',
                        f'attachment; filename= {filename}'
                    )
                    msg.attach(part)

        with smtplib.SMTP_SSL(smtp_config['server'], smtp_config['port']) as server:
            server.login(smtp_config['user'], smtp_config['password'])
            server.send_message(msg)
        return True, None
    except Exception as e:
        return False, str(e)

def extract_html_from_word(file_content):
    try:
        # mammoth требует путь к файлу, поэтому сохраняем content временно
        temp_file = "temp_template.docx"
        with open(temp_file, "wb") as f:
            f.write(file_content)
        with open(temp_file, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
        os.remove(temp_file)
        return html
    except Exception as e:
        print(f"❌ Ошибка при извлечении HTML из Word: {str(e)}")
        return None

def send_emails_view(request):
    processed_docx_path = request.session.get('processed_docx_path')
    last_uploaded_csv_path = request.session.get('last_uploaded_csv_path')
    last_processed_excel_path = request.session.get('last_processed_excel_path')

    print(f"DEBUG: processed_docx_path (before os.path.exists): {processed_docx_path}")
    print(f"DEBUG: os.path.exists(processed_docx_path): {os.path.exists(processed_docx_path)}")
    print(f"DEBUG: last_uploaded_csv_path (before os.path.exists): {last_uploaded_csv_path}")
    print(f"DEBUG: os.path.exists(last_uploaded_csv_path): {os.path.exists(last_uploaded_csv_path)}")

    docx_ready = processed_docx_path and os.path.exists(processed_docx_path)
    csv_ready = last_uploaded_csv_path and os.path.exists(last_uploaded_csv_path)
    print(f"DEBUG: docx_ready: {docx_ready}")
    print(f"DEBUG: csv_ready: {csv_ready}")

    if not docx_ready or not csv_ready:
        messages.error(request,
                       "Для рассылки необходимо сначала сгенерировать DOCX письмо и загрузить список рассылки.")
        return redirect('upload_excel')

    smtp_form = SmtpCredentialsForm()

    if request.method == 'POST':
        smtp_form = SmtpCredentialsForm(request.POST)
        if smtp_form.is_valid():
            email_account_value = smtp_form.cleaned_data['email_account']
            smtp_config = {
                'server': 'mail.hosting.reg.ru',
                'port': 465,
                'user': email_account_value,
                'password': smtp_form.cleaned_data['smtp_password'],
            }
            from_email = email_account_value
            subject_prefix = smtp_form.cleaned_data['subject_prefix']

            try:
                mailing_df = pd.read_csv(last_uploaded_csv_path)
                if 'Электронная почта' not in mailing_df.columns:
                    messages.error(request, "В CSV файле отсутствует колонка 'Электронная почта'.")
                    return render(request, 'main_app/send_emails.html', {
                                  'smtp_form': smtp_form, 'docx_ready': docx_ready, 'csv_ready': csv_ready,
                                  'project_name': request.session.get('project_name', 'Не определено'),
                                  'processed_docx_path': processed_docx_path,
                                  'last_processed_excel_path': last_processed_excel_path,
                                  'display_docx_filename': os.path.basename(processed_docx_path) if processed_docx_path else "(недоступно)",
                                  'display_excel_filename': os.path.basename(last_processed_excel_path) if last_processed_excel_path else "(недоступно)"})


                with open(processed_docx_path, 'rb') as f:
                    docx_content = f.read()
                html_body = extract_html_from_word(docx_content)
                if not html_body:
                    messages.error(request, "Не удалось извлечь содержимое письма из DOCX файла.")
                    return render(request, 'main_app/send_emails.html', {
                                  'smtp_form': smtp_form, 'docx_ready': docx_ready, 'csv_ready': csv_ready,
                                  'project_name': request.session.get('project_name', 'Не определено'),
                                  'processed_docx_path': processed_docx_path,
                                  'last_processed_excel_path': last_processed_excel_path,
                                  'display_docx_filename': os.path.basename(processed_docx_path) if processed_docx_path else "(недоступно)",
                                  'display_excel_filename': os.path.basename(last_processed_excel_path) if last_processed_excel_path else "(недоступно)"})

                sent_count = 0
                failed_recipients = []

                attachment_paths_for_email = []
                if last_processed_excel_path and os.path.exists(last_processed_excel_path):
                    attachment_paths_for_email.append(last_processed_excel_path)

                for index, row in mailing_df.iterrows():
                    to_emails = parse_emails(row.get('Электронная почта', ''))
                    if not to_emails:
                        messages.warning(request, f"Строка {index + 1}: Нет валидных email адресов для '{row.get('Наименование компании', 'Неизвестная компания')}'.")
                        continue

                    current_subject = f"{subject_prefix} - {request.session.get('project_name', 'Запрос КП')}"
                    email_body_with_placeholders = replace_placeholders(html_body, row)

                    success, error_msg = send_email_batch(
                        smtp_config, from_email, to_emails, current_subject, email_body_with_placeholders,
                        attachment_paths=attachment_paths_for_email
                    )

                    if success:
                        sent_count += 1
                    else:
                        messages.error(request, f"Отправка на {', '.join(to_emails)} не удалась: {error_msg}")
                        failed_recipients.append(f"Отправка на {', '.join(to_emails)} не удалась: {error_msg}")
                    time.sleep(1) # Задержка между отправками

                if sent_count > 0:
                    messages.success(request, f"Успешно отправлено {sent_count} писем.")
                if failed_recipients:
                    messages.warning(request, f"Не удалось отправить письма для: {len(failed_recipients)} получателей.")

                return redirect('send_emails')

            except FileNotFoundError as fnfe:
                messages.error(request, f"Ошибка: Отсутствует файл для рассылки ({fnfe}).")
            except Exception as e:
                messages.error(request, f"Ошибка при подготовке или отправке писем: {str(e)}")

        else:
            print("DEBUG: SMTP Form is NOT valid!")
            print("DEBUG: Form errors:", smtp_form.errors.as_json())
            messages.error(request, "Пожалуйста, исправьте ошибки в форме настроек SMTP.")

    display_docx_filename = "(недоступно)"
    if processed_docx_path:
        display_docx_filename = os.path.basename(processed_docx_path)

    display_excel_filename = "(недоступно)"
    if last_processed_excel_path:
        display_excel_filename = os.path.basename(last_processed_excel_path)

    return render(request, 'main_app/send_emails.html', {
        'smtp_form': smtp_form,
        'docx_ready': docx_ready,
        'csv_ready': csv_ready,
        'project_name': request.session.get('project_name', 'Не определено'),
        'processed_docx_path': processed_docx_path,
        'display_docx_filename': display_docx_filename,
        'display_excel_filename': display_excel_filename,
        'last_processed_excel_path': last_processed_excel_path,
    })
